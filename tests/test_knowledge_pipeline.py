from io import BytesIO

import pytest
from docx import Document

from app.services.docx_parser import DocumentBlock, DocxParser
from app.services.embedding_service import HybridEmbedding, SparseEmbedding, build_embedding_text
from app.services.vector_service import QdrantService, VectorPoint
from app.workflows.knowledge_ingestion import KnowledgeIngestionWorkflow


def test_docx_parser_preserves_paragraph_table_paragraph_order() -> None:
    document = Document()
    document.add_paragraph("表格前文字")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "条件"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "高开"
    table.cell(1, 1).text = "3%-7%"
    document.add_paragraph("表格后文字")
    stream = BytesIO()
    document.save(stream)

    parsed = DocxParser().parse(stream.getvalue())

    assert [block.type for block in parsed.blocks] == ["paragraph", "table", "paragraph"]
    assert parsed.blocks[0].text == "表格前文字"
    assert "| 条件 | 要求 |" in parsed.blocks[1].text
    assert parsed.blocks[2].text == "表格后文字"


def test_semantic_ranges_are_complete_and_non_overlapping() -> None:
    candidates = [
        {"start_block": 2, "end_block": 4, "title": "中间"},
        {"start_block": 4, "end_block": 99, "title": "末尾"},
    ]

    ranges = KnowledgeIngestionWorkflow._normalize_ranges(candidates, 7)
    indexes = [
        index
        for item in ranges
        for index in range(item["start_block"], item["end_block"] + 1)
    ]

    assert indexes == list(range(7))
    assert len(indexes) == len(set(indexes))


def test_max_length_protection_has_no_overlap() -> None:
    original = "投" * 6001
    groups = KnowledgeIngestionWorkflow._protect_max_length(
        [DocumentBlock(type="paragraph", index=0, text=original)]
    )

    contents = [block.content() for group in groups for block in group]
    assert "".join(contents) == original
    assert all(len(content) <= 6000 for content in contents)


def test_qdrant_rejects_authoritative_content_payload() -> None:
    service = QdrantService(
        url="http://127.0.0.1:6333",
        collection_name="test",
        dimension=2,
    )

    with pytest.raises(ValueError, match="禁止保存正文"):
        service.upsert(
            [
                VectorPoint(
                    id=1,
                    embedding=HybridEmbedding(
                        dense=[0.0, 1.0],
                        sparse=SparseEmbedding(indices=[1], values=[0.5]),
                    ),
                    payload={"chunk_id": 1, "content": "禁止进入 Qdrant"},
                )
            ]
        )


def test_embedding_text_is_context_newline_content() -> None:
    assert build_embedding_text("  N字战法卖点案例  ", "  原始正文  ") == "N字战法卖点案例\n原始正文"
    assert build_embedding_text(None, "  原始正文  ") == "原始正文"
    assert build_embedding_text("", "正文") == "正文"


def test_qdrant_allows_light_title_and_context_payload() -> None:
    QdrantService._validate_payload(
        {"chunk_id": 1, "title": "止盈处理", "context": "N字战法卖点案例"}
    )
